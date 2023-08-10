
import os,re
from translator.progress import Progress
from translator.llm_translator import LlmTranslator
from model import Model
from utils import LOG
from pdf2docx import Converter
from docx.shared import Cm
from docx import Document
# from docx2pdf import convert
from docx.enum.table import WD_ROW_HEIGHT_RULE

class DocParser:
    def __init__(self, model: Model, progress: Progress, target_language:str):
        self.translator = LlmTranslator(model, target_language)
        self.progress = progress
        self.taskConut = 0
        self.processed_paragraphs = set()  # 添加这行代码来存储已经处理过的段落

    def contains_email(text):
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
        # 使用正则表达式模式匹配邮箱地址
        if re.search(email_pattern, text):
            return True
        else:
            return False
    def is_float(self,text):
        try:
            float(text)
            return True
        except ValueError:
            return False

    def process_text(self, text, is_count_mode: bool):
        if(is_count_mode):
            self.taskConut+=1
            return
        else:
            self.progress.addCur()
        
        if text in ['\n', '']:
            return text
        if text.isnumeric() or self.is_float(text):
            return text
        translated_text = self.translator.translate_content(text)
        return translated_text
    
    
    def process_paragraph(self,paragraph, is_count_mode: bool):
            paragraph_hash = hash(paragraph.text)
            if paragraph_hash in self.processed_paragraphs:  # 更改为检查段落的哈希值
                return
            self.processed_paragraphs.add(paragraph_hash)
            # 检查段落中的每一个运行元素
            for run in paragraph.runs:
                # 如果运行元素包含图片或数学公式，跳过这个段落
                if 'graphicData' in run._r.xml or 'oMathPara' in run._r.xml:
                    return

            # 处理段落的文本
            processed_text = self.process_text(paragraph.text,is_count_mode)
            if is_count_mode:
                return
            
            print("翻译结果", processed_text)
            # 清空原有的段落
            paragraph.clear()
            # 添加处理后的文本到原有段落
            paragraph.add_run(processed_text)
    
    def process_all_paragraph(self, doc, is_count_mode: bool):
         # 遍历所有的段落
        if(doc.paragraphs is not None):
            for paragraph in doc.paragraphs:
                self.process_paragraph(paragraph,is_count_mode)
        return doc
    
    
    def process_table(self, table, is_count_mode: bool):
        # 遍历表格中的所有行
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO

            # 遍历行中的所有单元格
            for cell in row.cells:
                # 遍历单元格中的所有段落
                for paragraph in cell.paragraphs:
                    self.process_paragraph(paragraph, is_count_mode)
                    # paragraph.paragraph_format.space_after = Cm(0)  # 移除段落后的间距

                # 检查嵌套表格是否存在
                if cell.tables is not None:
                    # 递归处理嵌套表格
                    for nested_table in cell.tables:
                        self.process_table(nested_table, is_count_mode)
    
    def process_tables(self, doc, is_count_mode: bool):
        for table in doc.tables:
            self.process_table(table, is_count_mode)
                    

    def convertDoc(self,pdf_input_path,convert_docx_file_path,endPos):
        cv = Converter(pdf_input_path)
        # 执行转换操作
        cv.convert(convert_docx_file_path, start=0, end= endPos)
        # 关闭转换器，释放资源
        cv.close()

    def doTrans(self, file_name:str, result_docx_file_path:str, pdf_input_path: str, temp_source_path:str, endPos):
        endPos = endPos if endPos!=0 else None
        # 创建转换对象
        convert_docx_file_path = os.path.join(temp_source_path, file_name+".docx")
        self.convertDoc(pdf_input_path, convert_docx_file_path,endPos)
        
        self.taskConut = 0
        self.progress.resetCur()
        # 打开转换文档
        doc = Document(convert_docx_file_path)
        
        
        self.process_tables(doc, True)
        self.process_all_paragraph(doc, True)
        self.progress.setAll(self.taskConut)

        self.processed_paragraphs = set()  # 添加这行代码来存储已经处理过的段落
        self.process_tables(doc, False)
        self.process_all_paragraph(doc, False)
        
        # 保存修改后的文档
        doc.save(result_docx_file_path)
        print('处理完成！')



   

    