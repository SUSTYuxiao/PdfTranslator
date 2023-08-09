
from pdf2docx import Converter

# PDF文件路径
# pdf_file = "Attention Is All You Need.pdf"
pdf_file = "temp.pdf"

# 输出的DOCX文件路径
docx_file = "output4.docx"

# 创建转换器对象
cv = Converter(pdf_file)

# 执行转换操作
cv.convert(docx_file, start=0, end=None)

# 关闭转换器，释放资源
cv.close()
from docx import Document

def process_text(text):
    # 在这个函数中处理文本
    # 例如，我们可以将所有的文本转换为大写
    print(111)
    print(text)
    return text


def process_paragraph(paragraph):
    # 检查段落中的每一个运行元素
    for run in paragraph.runs:
        # 如果运行元素包含图片或数学公式，跳过这个段落
        if 'graphicData' in run._r.xml or 'oMathPara' in run._r.xml:
            print("jump", paragraph.text)
            return

    # 处理段落的文本
    processed_text = process_text(paragraph.text)

    # 清空原有的段落
    paragraph.clear()

    # 添加处理后的文本到原有段落
    paragraph.add_run(processed_text)

# 打开文档
doc = Document('output4.docx')


# 遍历所有的段落
for paragraph in doc.paragraphs:
    # 检查段落中的每一个运行元素
    for run in paragraph.runs:
        # 如果运行元素包含图片或数学公式，跳过这个段落
        if 'graphicData' in run._r.xml or 'oMathPara' in run._r.xml:
            break
    else:  # 没有找到图片或数学公式
        # 处理段落的文本
        processed_text = process_text(paragraph.text)

        # 清空原有的段落
        paragraph.clear()

        # 添加处理后的文本到原有段落
        paragraph.add_run(processed_text)
# 遍历所有的表格
for table in doc.tables:
    # 遍历表格中的所有行
    for row in table.rows:
        # 遍历行中的所有单元格
        for cell in row.cells:
            # 遍历单元格中的所有段落
            for paragraph in cell.paragraphs:
                print(333)
                print(paragraph.style.name, paragraph.text)
                process_paragraph(paragraph)

            # 遍历单元格中的所有表格
            for nested_table in cell.tables:
                # 遍历嵌套表格中的所有行
                for nested_row in nested_table.rows:
                    # 遍历行中的所有单元格
                    for nested_cell in nested_row.cells:
                        # 遍历单元格中的所有段落
                        print(444)
                        print(paragraph.style.name, paragraph.text)
                        for nested_paragraph in nested_cell.paragraphs:
                            process_paragraph(nested_paragraph)
